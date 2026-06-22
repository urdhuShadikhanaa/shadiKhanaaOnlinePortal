#!/usr/bin/env python3
"""Generate Urdu Shadikhana E2E Functionality Word document (styled)."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SCREENSHOTS = ROOT / "docs" / "screenshots"
OUTPUT = ROOT / "docs" / "Urdu-Shadikhana-E2E-Functionality-Reference.docx"

# Brand palette
PRIMARY = RGBColor(0x1B, 0x5E, 0x20)      # deep green
PRIMARY_LIGHT = RGBColor(0x2E, 0x7D, 0x32)
ACCENT = RGBColor(0xC9, 0xA2, 0x27)       # gold
HEADER_BG = "E8F5E9"
ALT_ROW_BG = "F5F5F5"
CALLOUT_BG = "FFF8E1"
FLOW_BG = "ECEFF1"
PLACEHOLDER_BG = "FAFAFA"

_bookmark_id = 0


class Theme:
    body_font = "Calibri"
    heading_font = "Calibri Light"
    mono_font = "Consolas"


def next_bookmark_id() -> int:
    global _bookmark_id
    _bookmark_id += 1
    return _bookmark_id


def setup_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = Theme.body_font
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    for level, size, color in [(1, 20, PRIMARY), (2, 15, PRIMARY_LIGHT), (3, 12, PRIMARY_LIGHT)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = Theme.heading_font
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.keep_with_next = True


def shade_cell(cell, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def add_bookmark(paragraph, bookmark_name: str) -> None:
    bid = next_bookmark_id()
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_hyperlink(paragraph, text: str, bookmark_name: str, bold: bool = False) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1B5E20")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(u)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_number_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run("Urdu Shadikhana — E2E Functionality Reference  |  Page ")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
        fld_run = p.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "separate")
        fld_char3 = OxmlElement("w:fldChar")
        fld_char3.set(qn("w:fldCharType"), "end")
        fld_run._r.append(fld_char1)
        fld_run._r.append(instr)
        fld_run._r.append(fld_char2)
        fld_run._r.append(fld_char3)


def add_header(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Urdu Shadikhana Online Reservation Portal")
        run.font.size = Pt(9)
        run.font.color.rgb = PRIMARY_LIGHT
        run.italic = True


def add_cover_page(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = band.rows[0].cells[0]
    shade_cell(cell, HEADER_BG)
    set_cell_margins(cell, 200, 200, 300, 300)

    t = cell.paragraphs[0]
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("URDU SHADIKHANA")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = PRIMARY
    r.font.name = Theme.heading_font

    sub = cell.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Online Reservation Portal")
    sr.font.size = Pt(18)
    sr.font.color.rgb = PRIMARY_LIGHT

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("End-to-End Functionality Reference")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = line.add_run("─" * 42)
    lr.font.color.rgb = ACCENT

    meta_tbl = doc.add_table(rows=4, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Document version", "1.1"),
        ("Generated", date.today().strftime("%d %B %Y")),
        ("Platform", "Salesforce Experience Cloud"),
        ("Integrations", "Email · SMS (Twilio / Android Gateway)"),
    ]
    for i, (k, v) in enumerate(meta_data):
        c0, c1 = meta_tbl.rows[i].cells
        c0.text = k
        c1.text = v
        for p in c0.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        for p in c1.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
        shade_cell(c0, ALT_ROW_BG)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("Confidential — For internal reference and future maintenance")
    nr.italic = True
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def add_word_toc_field(doc: Document) -> None:
    """Insert Word-native TOC field (user: References → Update Table)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)

    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char2)

    placeholder = doc.add_paragraph()
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = placeholder.add_run("[ Table of contents — open in Word and press Ctrl+A then F9 to update ]")
    pr.italic = True
    pr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    pr.font.size = Pt(10)

    run2 = p.add_run()
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run2._r.append(fld_char3)


def add_manual_toc(doc: Document, entries: list[tuple[int, str, str]]) -> None:
    """Styled clickable navigation TOC (works immediately in Word)."""
    toc_table = doc.add_table(rows=1, cols=2)
    toc_table.style = "Table Grid"
    hdr = toc_table.rows[0].cells
    hdr[0].text = "Section"
    hdr[1].text = "Title"
    for c in hdr:
        shade_cell(c, HEADER_BG)
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = PRIMARY

    for level, title, bookmark in entries:
        row = toc_table.add_row().cells
        sec = row[0].paragraphs[0]
        sec.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sec.add_run(title.split(".", 1)[0] if title[0].isdigit() else "—")
        sr.bold = True
        sr.font.size = Pt(10)
        sr.font.color.rgb = PRIMARY_LIGHT

        title_p = row[1].paragraphs[0]
        indent = "    " * (level - 1)
        add_hyperlink(title_p, f"{indent}{title}", bookmark, bold=(level == 1))
        if level > 1:
            for _ in range(level - 1):
                title_p.paragraph_format.left_indent = Inches(0.15 * (level - 1))

    doc.add_paragraph()
    tip = doc.add_paragraph()
    tip.paragraph_format.left_indent = Inches(0.2)
    tr = tip.add_run(
        "Tip: Click any section title above to jump to that section. "
        "For page numbers, update the automatic table below (Ctrl+A → F9 in Word)."
    )
    tr.italic = True
    tr.font.size = Pt(9)
    tr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def heading(doc: Document, text: str, level: int = 1, bookmark: str | None = None) -> None:
    bm = bookmark or f"heading_{next_bookmark_id()}"
    p = doc.add_heading(text, level=level)
    add_bookmark(p, bm)


def para(doc: Document, text: str, bold: bool = False, align=None) -> None:
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold


def callout(doc: Document, title: str, body: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, CALLOUT_BG)
    set_cell_margins(cell)
    p0 = cell.paragraphs[0]
    r0 = p0.add_run(title)
    r0.bold = True
    r0.font.color.rgb = PRIMARY
    p1 = cell.add_paragraph()
    r1 = p1.add_run(body)
    r1.font.size = Pt(10)
    doc.add_paragraph()


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(3)


def numbered_step(doc: Document, num: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    n = p.add_run(f"{num}. ")
    n.bold = True
    n.font.color.rgb = PRIMARY_LIGHT
    p.add_run(text)


def flow_box(doc: Document, lines: list[str], title: str | None = None) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, FLOW_BG)
    set_cell_margins(cell)
    if title:
        pt = cell.paragraphs[0]
        tr = pt.add_run(title)
        tr.bold = True
        tr.font.size = Pt(10)
        tr.font.color.rgb = PRIMARY
    body = cell.add_paragraph() if title else cell.paragraphs[0]
    run = body.add_run("\n".join(lines))
    run.font.name = Theme.mono_font
    run.font.size = Pt(8.5)
    doc.add_paragraph()


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], HEADER_BG)
        set_cell_margins(hdr[i])
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = PRIMARY
        r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_margins(cells[i])
            if ri % 2 == 1:
                shade_cell(cells[i], ALT_ROW_BG)
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.size = Pt(10)
    doc.add_paragraph()


def screenshot(doc: Document, filename: str, caption: str, placeholder: str | None = None) -> None:
    frame = doc.add_table(rows=1, cols=1)
    frame.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = frame.rows[0].cells[0]
    shade_cell(cell, PLACEHOLDER_BG if not (SCREENSHOTS / filename).exists() else "FFFFFF")
    set_cell_margins(cell, 120, 120, 120, 120)
    path = SCREENSHOTS / filename
    if path.exists():
        pic_p = cell.paragraphs[0]
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.add_run().add_picture(str(path), width=Inches(5.8))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    else:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("SCREENSHOT PLACEHOLDER")
        run.bold = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        hint = cell.add_paragraph()
        hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hint.add_run(placeholder or caption)
        hr.italic = True
        hr.font.size = Pt(9)
        hr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        fn = cell.add_paragraph()
        fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fn.add_run(f"File: docs/screenshots/{filename}")
        fr.font.name = Theme.mono_font
        fr.font.size = Pt(8)
    doc.add_paragraph()


def scenario_section_title(doc: Document, number: str, title: str) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.columns[0].width = Inches(0.9)
    c0, c1 = tbl.rows[0].cells
    shade_cell(c0, HEADER_BG)
    shade_cell(c1, "FFFFFF")
    set_cell_margins(c0)
    set_cell_margins(c1)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(number)
    r0.bold = True
    r0.font.size = Pt(14)
    r0.font.color.rgb = PRIMARY
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(14)
    r1.font.color.rgb = PRIMARY_LIGHT
    doc.add_paragraph()


def scenario(
    doc: Document,
    number: str,
    title: str,
    actor: str,
    purpose: str,
    steps: list[str],
    components: list[list[str]],
    screenshot_files: list[tuple[str, str, str | None]],
    notifications: str | None = None,
) -> None:
    bookmark = f"scenario_{number.lower()}"
    scenario_section_title(doc, number, title)
    anchor = doc.add_paragraph()
    anchor.paragraph_format.space_before = Pt(0)
    anchor.paragraph_format.space_after = Pt(4)
    add_bookmark(anchor, bookmark)

    table(
        doc,
        ["Attribute", "Description"],
        [
            ["Actor", actor],
            ["Purpose", purpose],
            ["Notifications", notifications or "None"],
        ],
    )

    para(doc, "Process flow", bold=True)
    for i, step in enumerate(steps, 1):
        numbered_step(doc, i, step)

    para(doc, "Related components", bold=True)
    table(doc, ["Layer", "Component", "Path / Role"], components)

    para(doc, "User interface", bold=True)
    for fname, cap, hint in screenshot_files:
        screenshot(doc, fname, cap, hint)

    doc.add_paragraph()  # spacing between scenarios


def build_document() -> Document:
    global _bookmark_id
    _bookmark_id = 0

    doc = Document()
    setup_document(doc)
    add_header(doc)
    add_page_number_footer(doc)
    add_cover_page(doc)

    # ── Table of Contents page ──
    heading(doc, "Table of Contents", level=1, bookmark="sec_toc")
    callout(
        doc,
        "How to use this document",
        "Section 2 contains master flow diagrams. Section 3 walks through every user scenario "
        "with screenshots and component lists. Section 4 is a complete technical reference. "
        "Use the navigation table below or update the automatic table for page numbers.",
    )
    heading(doc, "Quick navigation", level=2)
    # TOC entries collected during build — we need two-pass OR pre-register main sections
    toc_entries = [
        (1, "1. Solution Overview", "sec_pre_1"),
        (1, "2. Master Flow Diagrams", "sec_pre_2"),
        (1, "3. Functional Scenarios", "sec_pre_3"),
        (1, "4. Complete Component Reference", "sec_pre_4"),
        (1, "5. Screenshot Capture Guide", "sec_pre_5"),
        (1, "6. Related Documentation", "sec_pre_6"),
    ] + [
        (2, f"Scenario {s}: {t}", f"scenario_{s.lower()}")
        for s, t in [
            ("S01", "Portal Home & Header"),
            ("S02", "Venue Gallery & Address"),
            ("S03", "Booking Wizard — Step 1"),
            ("S04", "Booking Wizard — Step 2"),
            ("S05", "Booking Wizard — Step 3"),
            ("S06", "Booking Wizard — Step 4"),
            ("S07", "Notifications on New Booking"),
            ("S08", "Experience Cloud Login"),
            ("S09", "Admin — Booking Queue"),
            ("S10", "Admin — Booking Report"),
            ("S11", "Admin — Pricing"),
            ("S12", "Admin — Site Branding"),
            ("S13", "Admin — SMS Settings"),
            ("S14", "Admin — Email Settings"),
            ("S15", "Admin — Portal Login Notes"),
            ("S16", "Admin — Booking Data"),
            ("S17", "Android SMS Gateway"),
            ("S18", "Scheduled Cleanup"),
        ]
    ]
    add_manual_toc(doc, toc_entries)
    heading(doc, "Automatic table (with page numbers)", level=2)
    add_word_toc_field(doc)
    doc.add_page_break()

    # ── 1 Overview ──
    heading(doc, "1. Solution Overview", level=1, bookmark="sec_pre_1")
    callout(
        doc,
        "At a glance",
        "Urdu Shadikhana is a Salesforce Experience Cloud portal for a community event hall in "
        "Nandalur, Andhra Pradesh. Customers submit booking requests online; administrators review, "
        "confirm or cancel, manage pricing, and configure email/SMS notifications.",
    )
    heading(doc, "1.1 User Roles", level=2)
    table(
        doc,
        ["Role", "Portal access", "Sign-in method"],
        [
            ["Guest (public)", "Home, venue info, 4-step booking wizard", "No login required"],
            ["Community user", "Same as guest (Home tab)", "Experience Cloud login"],
            ["System Administrator", "Home + Booking Data + Administration (7 sub-sections)", "Salesforce System Administrator profile"],
        ],
    )
    heading(doc, "1.2 Technology Stack", level=2)
    table(
        doc,
        ["Layer", "Technology"],
        [
            ["Portal UI", "Experience Cloud (Urdu_Shadikhana1) + LWC shadikhanaBookingPortal"],
            ["Backend", "Apex — ShadikhanaBookingController + service classes"],
            ["Data model", "Shadikhana_Booking__c, Shadikhana_Settings__c, Sms_Outbound__c"],
            ["Email", "Messaging.sendEmail via ShadikhanaBookingEmailService"],
            ["SMS", "Twilio callout OR Android Gateway (Sms_Outbound__c queue)"],
            ["Mobile gateway", "Kotlin app — android/shadikhana-sms-gateway"],
        ],
    )

    doc.add_page_break()

    # ── 2 Flow diagrams ──
    heading(doc, "2. Master Flow Diagrams", level=1, bookmark="sec_pre_2")

    heading(doc, "2.1 System Architecture", level=2)
    flow_box(
        doc,
        [
            "┌─────────────────┐     HTTPS      ┌──────────────────────────────┐",
            "│  Web Browser    │ ─────────────► │  Experience Cloud Portal     │",
            "│  (Guest/Admin)  │                │  c:shadikhanaBookingPortal   │",
            "└─────────────────┘                └──────────────┬───────────────┘",
            "                                                  │ @AuraEnabled",
            "                                                  ▼",
            "                                   ┌──────────────────────────────┐",
            "                                   │  ShadikhanaBookingController │",
            "                                   └──────────────┬───────────────┘",
            "                    ┌────────────────────────────┼────────────────────────────┐",
            "                    ▼                            ▼                            ▼",
            "         ShadikhanaBookingPricing    ShadikhanaBookingEmailService   ShadikhanaBookingSmsService",
            "                    │                            │                            │",
            "                    ▼                            ▼                            ▼",
            "         Shadikhana_Booking__c          Email (requester/admin)     Twilio OR Sms_Outbound__c",
            "         Shadikhana_Daily_Rate__c                                              │",
            "                                                                               ▼",
            "                                                                   ┌──────────────────┐",
            "                                                                   │ Android SMS App  │",
            "                                                                   └──────────────────┘",
        ],
        title="Figure 2.1 — High-level architecture",
    )

    heading(doc, "2.2 Booking Lifecycle", level=2)
    flow_box(
        doc,
        [
            "  [Guest submits form]",
            "         │",
            "         ▼",
            "    ┌──────────┐",
            "    │ Pending  │ ◄── Admin + Requester: SMS & Email notifications",
            "    └────┬─────┘",
            "         │ Admin confirms (no overlap with other Confirmed)",
            "         ▼",
            "    ┌───────────┐",
            "    │ Confirmed │ ◄── Requester SMS + Email; calendar blocked",
            "    └─────┬─────┘",
            "          │ Admin cancels",
            "          ▼",
            "    ┌───────────┐",
            "    │ Cancelled │ ◄── Requester SMS + Email; Final Amount = ₹0",
            "    └───────────┘",
        ],
        title="Figure 2.2 — Booking status lifecycle",
    )

    heading(doc, "2.3 Notification Matrix", level=2)
    table(
        doc,
        ["Event", "Email → Requester", "Email → Admin", "SMS → Requester", "SMS → Admin"],
        [
            ["New booking", "Yes (if email provided)", "Yes + review link", "Yes — pending ack", "Yes + review link"],
            ["Confirmed", "Yes", "No", "Yes", "No"],
            ["Cancelled", "Yes", "No", "Yes", "No"],
        ],
    )

    heading(doc, "2.4 Android SMS Gateway", level=2)
    flow_box(
        doc,
        [
            "Apex: ShadikhanaBookingSmsService → ShadikhanaSmsGatewayService.enqueueMessage()",
            "  → INSERT Sms_Outbound__c (Status = Pending)",
            "",
            "Android app polls every ~30–60 seconds:",
            "  GET  /services/apexrest/shadikhana/sms/v1/pending",
            "  → Send SMS via device SIM",
            "  POST /services/apexrest/shadikhana/sms/v1/ack",
            "  → UPDATE Sms_Outbound__c (Sent | Failed)",
        ],
        title="Figure 2.4 — Android gateway message flow",
    )

    heading(doc, "2.5 Portal Navigation", level=2)
    flow_box(
        doc,
        [
            "Sidebar navigation (LWC client-side)",
            "├── Home ........................ Everyone",
            "├── Booking Data ................ System Administrator only",
            "└── Administration",
            "    ├── Booking Queue",
            "    ├── Booking Report",
            "    ├── Pricing",
            "    ├── Site Branding",
            "    ├── SMS Notifications",
            "    ├── Email Notifications",
            "    └── Portal Login Notes",
        ],
        title="Figure 2.5 — Portal navigation map",
    )

    doc.add_page_break()

    # ── 3 Scenarios ──
    heading(doc, "3. Functional Scenarios", level=1, bookmark="sec_pre_3")
    para(
        doc,
        "Each scenario documents the end-user flow, Salesforce components involved, and UI screenshots. "
        "Placeholder images can be replaced under docs/screenshots/ and the document regenerated.",
    )

    scenario(
        doc, "S01", "Portal Home & Header", "Guest / All users",
        "View venue branding, login button, and navigation sidebar.",
        [
            "User opens Experience Cloud URL (/urdushadikhana).",
            "home.json loads c:shadikhanaBookingPortal.",
            "getCurrentUserInfo determines guest vs admin.",
            "Header shows venue name, clock, and Login button (guests).",
        ],
        [
            ["Experience", "Urdu_Shadikhana1 / home.json", "experiences/Urdu_Shadikhana1/views/home.json"],
            ["LWC", "shadikhanaBookingPortal", "lwc/shadikhanaBookingPortal/"],
            ["Apex", "getCurrentUserInfo", "ShadikhanaBookingController.cls"],
        ],
        [
            ("01-portal-home-header.png", "Figure S01a — Portal header and sidebar", None),
            ("02-portal-home-top.png", "Figure S01b — Home top section", None),
        ],
    )

    scenario(
        doc, "S02", "Venue Gallery & Address", "Guest / All users",
        "Review hall photos and venue location before booking.",
        [
            "Scroll to 'Review the Shadikhana' gallery.",
            "Navigate photos via thumbnails or arrows.",
            "View Venue Address card and Google Maps embed.",
        ],
        [
            ["LWC", "Gallery section", "shadikhanaBookingPortal.html/js"],
            ["Static Resource", "Hall images", "ShadikhanaHallExterior, Interior, DiningArea"],
            ["Apex", "getPortalBrandingConfig", "ShadikhanaPortalBrandingService"],
        ],
        [
            ("03-venue-gallery.png", "Figure S02a — Photo gallery", None),
            ("04-venue-address-map.png", "Figure S02b — Venue address and map", None),
        ],
    )

    scenario(
        doc, "S03", "Booking Wizard — Step 1 (Dates & Times)", "Guest / Community user",
        "Select available date range and event timing.",
        [
            "Open Booking Request Form — Step 1 of 4.",
            "Pick From Date and To Date on inline calendar.",
            "Unavailable: past dates, confirmed, or pending bookings.",
            "Set Start Time and End Time.",
            "Right panel shows availability calendar with colour legend.",
        ],
        [
            ["LWC", "validateBookingWizardStep1", "shadikhanaBookingPortal.js"],
            ["Apex", "getBookings, isDateRangeAvailable", "ShadikhanaBookingController"],
            ["Apex", "calculateBookingPrice", "ShadikhanaBookingPricing"],
            ["Object", "Shadikhana_Booking__c", "Booking_Date__c, End_Date__c, Status__c"],
        ],
        [
            ("05-booking-form-step1-dates.png", "Figure S03a — Booking form step 1", "Capture date pickers"),
            ("06-calendar-date-picker.png", "Figure S03b — Date selection", None),
            ("07-availability-calendar-panel.png", "Figure S03c — Availability panel", None),
        ],
    )

    scenario(
        doc, "S04", "Booking Wizard — Step 2 (Event Information)", "Guest / Community user",
        "Enter event type, name, guest count, decoration and catering.",
        [
            "Select Event Type from picklist.",
            "Enter Event Name and Expected Guest Count.",
            "Choose Decoration and Catering arrangements.",
            "Optional: Event Details textarea.",
            "Estimated price updates automatically.",
        ],
        [
            ["LWC", "validateBookingWizardStep2", "shadikhanaBookingPortal.js"],
            ["Apex", "calculateBookingPrice", "ShadikhanaBookingPricing"],
            ["Settings", "Hall_* charge fields", "Shadikhana_Settings__c"],
        ],
        [("10-booking-step2-event.png", "Figure S04 — Event information step", "Advance to step 2")],
    )

    scenario(
        doc, "S05", "Booking Wizard — Step 3 (Contact Information)", "Guest / Community user",
        "Provide contact details for notifications.",
        [
            "Enter Contact Name (required).",
            "Enter 10-digit Mobile Number (required).",
            "Enter Email (optional — used for acknowledgment).",
            "Optional notes for admin team.",
        ],
        [
            ["LWC", "validateBookingWizardStep3", "Phone and name validation"],
            ["Apex", "createBooking contact fields", "Contact_Name__c, Contact_Phone__c, Contact_Email__c"],
        ],
        [("11-booking-step3-contact.png", "Figure S05 — Contact information", None)],
    )

    scenario(
        doc, "S06", "Booking Wizard — Step 4 (Review & Submit)", "Guest / Community user",
        "Review summary, accept consent, and submit booking request.",
        [
            "Review booking preview and estimated price.",
            "Accept terms/consent checkbox.",
            "Click Submit Booking Request.",
            "createBooking inserts record with Status = Pending.",
            "Success toast displayed; calendar refreshes.",
        ],
        [
            ["LWC", "handleSubmitBooking", "shadikhanaBookingPortal.js"],
            ["Apex", "createBooking", "ShadikhanaBookingController.cls"],
            ["Object", "Shadikhana_Booking__c", "All pricing and status fields"],
        ],
        [
            ("12-booking-step4-review.png", "Figure S06a — Review step", None),
            ("13-booking-submit-success.png", "Figure S06b — Success confirmation", None),
        ],
        notifications="Email + Admin SMS + Requester SMS (see S07)",
    )

    scenario(
        doc, "S07", "Notifications on New Booking", "System (automated)",
        "Alert admin and acknowledge requester after booking creation.",
        [
            "notifyBookingRequestSubmitted queues requester + admin emails.",
            "notifyAdminNewBooking sends SMS to all admin mobile numbers.",
            "notifyRequesterBookingSubmitted sends acknowledgment SMS.",
            "Android path: Sms_Outbound__c Pending → app polls → Sent.",
        ],
        [
            ["Apex", "ShadikhanaBookingEmailService", "BookingRequestEmailQueueable"],
            ["Apex", "ShadikhanaBookingSmsService", "Admin + Requester queueables"],
            ["Apex", "ShadikhanaSmsGatewayRest", "REST /pending and /ack"],
            ["Object", "Sms_Outbound__c", "SMS queue for Android gateway"],
            ["Android", "SmsGatewayService", "android/shadikhana-sms-gateway/"],
        ],
        [
            ("14-sms-outbound-pending.png", "Figure S07a — Sms_Outbound__c records", "Salesforce list view"),
            ("15-android-gateway-app.png", "Figure S07b — Android gateway app", "Phone screenshot"),
            ("16-admin-sms-received.png", "Figure S07c — Admin SMS with link", "Phone screenshot"),
        ],
        notifications="See Section 2.3 notification matrix",
    )

    scenario(
        doc, "S08", "Experience Cloud Login", "Community user",
        "Sign in to the portal.",
        [
            "Click Login in portal header.",
            "login.json shows logo, shadikhanaLoginNotes, and login form.",
            "User authenticates via Salesforce Identity.",
            "Redirected back to portal Home.",
        ],
        [
            ["Experience", "login.json", "experiences/Urdu_Shadikhana1/views/login.json"],
            ["LWC", "shadikhanaLoginNotes", "lwc/shadikhanaLoginNotes/"],
            ["Settings", "Portal_Login_Notes__c", "Shadikhana_Settings__c"],
        ],
        [("09-login-page.png", "Figure S08 — Login page", None)],
    )

    scenario(
        doc, "S09", "Admin — Booking Queue", "System Administrator",
        "Review pending requests; confirm, cancel, or set final price.",
        [
            "Open Administration → Booking Queue.",
            "getAdminBookings loads cards grouped by status.",
            "Confirm → overlap check + requester notifications.",
            "Cancel → Final_Amount__c = 0.",
            "Deep link ?bookingId=&tab=admin highlights booking card.",
        ],
        [
            ["LWC", "handleConfirmBooking, handleCancelBooking", "shadikhanaBookingPortal.js"],
            ["Apex", "updateBookingStatus, updateBookingPrice", "ShadikhanaBookingController"],
            ["Apex", "Email + SMS services", "Status-change notifications"],
        ],
        [
            ("20-admin-booking-queue.png", "Figure S09a — Booking Queue", "Admin login required"),
            ("21-admin-confirm-booking.png", "Figure S09b — Confirm action", None),
            ("22-admin-deep-link-highlight.png", "Figure S09c — Deep link highlight", None),
        ],
    )

    scenario(
        doc, "S10", "Admin — Booking Report", "System Administrator",
        "View monthly or custom-range booking analytics.",
        ["Open Administration → Booking Report.", "Select month/year or custom range.", "View summary and day-by-day breakdown."],
        [["Apex", "getMonthlyAdminReport", "ShadikhanaBookingController"]],
        [("23-admin-booking-report.png", "Figure S10 — Booking Report", None)],
    )

    scenario(
        doc, "S11", "Admin — Pricing Management", "System Administrator",
        "Configure default daily rate and per-date overrides.",
        ["Open Administration → Pricing.", "Edit default daily rate.", "Set per-date overrides on calendar."],
        [["Apex", "saveDateRateOverride", "ShadikhanaBookingController"], ["Object", "Shadikhana_Daily_Rate__c", "Per-date rates"]],
        [("24-admin-pricing.png", "Figure S11 — Pricing section", None)],
    )

    scenario(
        doc, "S12", "Admin — Site Branding", "System Administrator",
        "Update venue name, address, and gallery images.",
        ["Open Administration → Site Branding.", "Edit venue name and addresses.", "Upload gallery images."],
        [["Apex", "ShadikhanaPortalBrandingService", "Branding persistence"]],
        [("25-admin-branding.png", "Figure S12 — Site Branding", None)],
    )

    scenario(
        doc, "S13", "Admin — SMS Notifications Settings", "System Administrator",
        "Configure SMS provider and admin mobile numbers.",
        ["Open SMS Notifications section.", "Toggle SMS Enabled.", "Set admin mobiles; configure Twilio or Android Gateway."],
        [["LWC", "isAdminSmsSection", "shadikhanaBookingPortal"], ["Permission Set", "Shadikhana_Sms_Gateway", "Android REST access"]],
        [("26-admin-sms-settings.png", "Figure S13 — SMS settings", None)],
    )

    scenario(
        doc, "S14", "Admin — Email Notifications Settings", "System Administrator",
        "Configure email notifications and portal URLs.",
        ["Open Email Notifications section.", "Toggle Email Enabled; set Admin Email.", "Set Portal Site URL for deep links."],
        [["Apex", "ShadikhanaBookingEmailService", "Email queueables"]],
        [("27-admin-email-settings.png", "Figure S14 — Email settings", None)],
    )

    scenario(
        doc, "S15", "Admin — Portal Login Notes", "System Administrator",
        "Customize text on the Experience login page.",
        ["Open Portal Login Notes.", "Edit heading, message, credential hints.", "Save — reflected on login page."],
        [["LWC", "shadikhanaLoginNotes", "Login page component"]],
        [("28-admin-login-notes.png", "Figure S15 — Login notes editor", None)],
    )

    scenario(
        doc, "S16", "Admin — Booking Data Tab", "System Administrator",
        "Paginated read-only view of all booking records.",
        ["Open Booking Data sidebar tab.", "getBookingData (admin-only).", "Paginate 5–50 records per page."],
        [["Apex", "getBookingData", "ShadikhanaBookingController"]],
        [("29-admin-booking-data.png", "Figure S16 — Booking Data table", None)],
    )

    scenario(
        doc, "S17", "Android SMS Gateway — Setup & Operation", "Technical admin",
        "Send SMS via physical Android phone SIM.",
        [
            "Deploy Sms_Outbound__c, gateway Apex, permission set.",
            "Set SMS_Provider__c = Android Gateway in settings.",
            "Install APK; configure instance URL, OAuth token, API key.",
            "Start foreground polling service.",
        ],
        [
            ["Android", "MainActivity.kt, SmsGatewayService.kt", "android/shadikhana-sms-gateway/"],
            ["Apex", "ShadikhanaSmsGatewayRest", "REST endpoints"],
        ],
        [
            ("15-android-gateway-app.png", "Figure S17a — Gateway app", "Phone screenshot"),
            ("30-android-test-sms.png", "Figure S17b — Test SMS", "Phone screenshot"),
        ],
    )

    scenario(
        doc, "S18", "Scheduled Cancelled Booking Cleanup", "System (scheduled)",
        "Remove old cancelled bookings monthly.",
        ["Scheduler runs 5th of each month.", "Batch deletes prior month Cancelled bookings."],
        [["Apex", "ShadikhanaCancelledCleanupScheduler", "Scheduled job"]],
        [("31-scheduled-jobs.png", "Figure S18 — Scheduled Jobs", "Setup → Scheduled Jobs")],
    )

    doc.add_page_break()

    # ── 4 Reference ──
    heading(doc, "4. Complete Component Reference", level=1, bookmark="sec_pre_4")
    heading(doc, "4.1 Lightning Web Components", level=2)
    table(
        doc, ["Component", "Purpose", "Path"],
        [
            ["shadikhanaBookingPortal", "Main portal (all tabs, wizard, admin)", "lwc/shadikhanaBookingPortal/"],
            ["shadikhanaLoginNotes", "Login page notes", "lwc/shadikhanaLoginNotes/"],
            ["shadikhanaLoginBranding", "Login branding (optional)", "lwc/shadikhanaLoginBranding/"],
        ],
    )
    heading(doc, "4.2 Apex Classes", level=2)
    table(
        doc, ["Class", "Purpose"],
        [
            ["ShadikhanaBookingController", "All @AuraEnabled portal APIs"],
            ["ShadikhanaBookingPricing", "Price calculation and rate overrides"],
            ["ShadikhanaBookingEmailService", "Email notifications"],
            ["ShadikhanaBookingSmsService", "SMS (Twilio + gateway routing)"],
            ["ShadikhanaSmsGatewayService", "Sms_Outbound__c queue"],
            ["ShadikhanaSmsGatewayRest", "REST API for Android app"],
            ["ShadikhanaPortalBrandingService", "Branding and gallery"],
            ["ShadikhanaCancelledBookingCleanupBatch", "Monthly cleanup"],
        ],
    )
    heading(doc, "4.3 Custom Objects", level=2)
    table(
        doc, ["Object", "Purpose"],
        [
            ["Shadikhana_Booking__c", "Booking requests and pricing"],
            ["Shadikhana_Settings__c", "Org-wide portal and notification settings"],
            ["Shadikhana_Daily_Rate__c", "Per-date pricing overrides"],
            ["Sms_Outbound__c", "Android SMS outbound queue"],
            ["Portal_Banner__c", "Portal banner messages"],
        ],
    )
    heading(doc, "4.4 Permission Sets", level=2)
    table(
        doc, ["Permission Set", "Purpose"],
        [
            ["Shadikhana_Guest_Portal", "Public guest booking access"],
            ["Shadikhana_Community_Booking", "Logged-in community users"],
            ["Shadikhana_Booking_User", "Internal admin booking CRUD"],
            ["Shadikhana_Twilio_Callout", "Twilio named credential"],
            ["Shadikhana_Sms_Gateway", "Android integration user"],
        ],
    )

    doc.add_page_break()

    heading(doc, "5. Screenshot Capture Guide", level=1, bookmark="sec_pre_5")
    callout(
        doc,
        "Regenerating this document",
        "1. Add PNG files to docs/screenshots/\n"
        "2. Run: python scripts/generate-e2e-word-doc.py\n"
        "3. Open in Word and press Ctrl+A then F9 to refresh page numbers",
    )
    bullet(doc, "Public pages: node scripts/capture-doc-screenshots.js")
    bullet(doc, "Admin pages: capture manually while logged in as System Administrator")
    bullet(doc, "See docs/screenshots/README.md for filename list")

    heading(doc, "5.1 Screenshot Index", level=2)
    rows = []
    if SCREENSHOTS.exists():
        for f in sorted(SCREENSHOTS.glob("*.png")):
            rows.append([f.name, "✓ Captured", ""])
    if not rows:
        rows.append(["(none)", "—", "Run capture script"])
    table(doc, ["Filename", "Status", "Notes"], rows)

    heading(doc, "6. Related Documentation", level=1, bookmark="sec_pre_6")
    table(
        doc, ["Document", "Path", "Audience"],
        [
            ["Functionality Guide", "docs/FUNCTIONALITY-GUIDE.md", "Business users"],
            ["Technical Documentation", "docs/TECHNICAL-DOCUMENTATION.md", "Developers"],
            ["Deployment Guide", "docs/DEPLOYMENT-GUIDE.md", "DevOps"],
            ["Android Gateway README", "android/shadikhana-sms-gateway/README.md", "Mobile / SMS"],
            ["Post-Deploy Checklist", "deploy/urdu-shadikhana/POST-DEPLOY-CHECKLIST.md", "Deploy"],
        ],
    )

    return doc


def main() -> None:
    SCREENSHOTS.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(OUTPUT))
    print(f"Created: {OUTPUT}")
    print("Open in Microsoft Word and press Ctrl+A then F9 to update the table of contents page numbers.")


if __name__ == "__main__":
    main()
